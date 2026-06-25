import streamlit as st
import pandas as pd
from datetime import date
import os
from io import BytesIO
from reportlab.pdfgen import canvas
from docx import Document

FILE = "data.csv"

st.set_page_config(page_title="J&A Tequeños", layout="wide")

# 🎨 DESIGN PREMIUM
st.markdown("""
<style>
body {
    background: linear-gradient(180deg, #020617, #0F172A);
    color: #E2E8F0;
}

.logo {
    text-align:center;
    font-size:40px;
    font-weight:bold;
    background: linear-gradient(90deg,#22c55e,#3b82f6);
    -webkit-background-clip:text;
    -webkit-text-fill-color:transparent;
}

.banner {
    height:200px;
    background-image:url('https://domesticfits.com/wp-content/uploads/2023/06/venezuelan-food-Tequenos.jpg');
    background-size:cover;
    border-radius:15px;
    margin-bottom:20px;
}

.card {
    background: linear-gradient(145deg,#1E293B,#111827);
    padding:20px;
    border-radius:15px;
    text-align:center;
    transition:0.3s;
}
.card:hover {
    transform:scale(1.05);
}

.value {
    font-size:30px;
    font-weight:bold;
    color:white;
}

.green {color:#22c55e;}
.red {color:#ef4444;}
.blue {color:#3b82f6;}

.section {
    background:#1E293B;
    padding:20px;
    border-radius:15px;
    margin-bottom:20px;
}

.stButton>button {
    background:linear-gradient(90deg,#22c55e,#15803d);
    color:white;
    border-radius:10px;
}
</style>
""", unsafe_allow_html=True)

# ✅ HEADER
st.markdown("<div class='logo'>🌮 J&A Tequeños</div>", unsafe_allow_html=True)
st.markdown("<div class='banner'></div>", unsafe_allow_html=True)

PRICES = {"Tequeños":4,"Pasteles":5}

# ✅ LOAD DATA
if os.path.exists(FILE):
    df = pd.read_csv(FILE)
else:
    df = pd.DataFrame(columns=["tipo","producto","cantidad","precio_unitario","monto","fecha"])

df = df.dropna(how="all")

# ✅ IMPORT CSV (RESTORE)
st.markdown("<div class='section'>", unsafe_allow_html=True)
st.subheader("📂 Restaurar datos")

uploaded_file = st.file_uploader("Subir backup CSV")

if uploaded_file:
    df = pd.read_csv(uploaded_file)
    df.to_csv(FILE, index=False)
    st.success("✅ Datos restaurados correctamente")

st.markdown("</div>", unsafe_allow_html=True)

# ✅ FORM
st.markdown("<div class='section'>", unsafe_allow_html=True)
st.subheader("➕ Nueva transacción")

col1,col2 = st.columns(2)

tipo = col1.selectbox("Tipo",["Venta","Gasto"])

if tipo=="Venta":
    producto = col1.selectbox("Producto",["Tequeños","Pasteles"])
    cantidad = col2.number_input("Cantidad",min_value=1)

    precio = PRICES[producto]
    monto = cantidad * precio

    st.success(f"Ingreso: ${monto}")

else:
    producto = col1.text_input("Material")
    cantidad = col2.number_input("Cantidad",min_value=1)
    precio = col2.number_input("Precio",min_value=0.0)

    monto = cantidad * precio

    st.warning(f"Gasto: ${monto}")

fecha = st.date_input("Fecha", value=date.today())

if st.button("Guardar"):
    new = pd.DataFrame([{
        "tipo":tipo,
        "producto":producto,
        "cantidad":cantidad,
        "precio_unitario":precio,
        "monto":monto,
        "fecha":fecha
    }])
    df = pd.concat([df,new],ignore_index=True)
    df.to_csv(FILE,index=False)
    st.success("✅ Guardado correctamente")

st.markdown("</div>", unsafe_allow_html=True)

# ✅ DISPLAY
if len(df)>0:

    ingresos = df[df["tipo"]=="Venta"]["monto"].sum()
    gastos = df[df["tipo"]=="Gasto"]["monto"].sum()
    beneficio = ingresos - gastos

    st.subheader("📊 Resumen")

    c1,c2,c3 = st.columns(3)

    c1.markdown(f"""
    <div class='card'><div class='blue'>Ingresos</div>
    <div class='value'>${ingresos:.2f}</div></div>
    """,unsafe_allow_html=True)

    c2.markdown(f"""
    <div class='card'><div class='red'>Gastos</div>
    <div class='value'>${gastos:.2f}</div></div>
    """,unsafe_allow_html=True)

    c3.markdown(f"""
    <div class='card'><div class='green'>Beneficio</div>
    <div class='value'>${beneficio:.2f}</div></div>
    """,unsafe_allow_html=True)

    # ✅ GRAPH
    st.subheader("📈 Evolución")

    try:
        df_copy = df.copy()
        df_copy["fecha"] = pd.to_datetime(df_copy["fecha"])

        daily = df_copy.groupby(["fecha","tipo"])["monto"].sum().unstack(fill_value=0)

        if "Venta" not in daily: daily["Venta"]=0
        if "Gasto" not in daily: daily["Gasto"]=0

        daily["Beneficio"] = daily["Venta"] - daily["Gasto"]

        daily = daily.rename(columns={
            "Venta":"Ingresos",
            "Gasto":"Gastos"
        })

        st.line_chart(daily)

    except Exception as e:
        st.error(e)

    # ✅ HISTORIAL
    st.subheader("📋 Historial")
    st.dataframe(df, use_container_width=True)

    # ✅ DELETE ONE LINE
    st.subheader("🗑️ Eliminar entrada")

    idx = st.number_input("Número de fila", 0, len(df)-1, 0)

    if st.button("Eliminar"):
        df = df.drop(index=idx)
        df.to_csv(FILE, index=False)
        st.success("✅ Entrada eliminada")

    # ✅ EXPORT + BACKUP
    st.subheader("💾 Seguridad y exportación")

    st.warning("⚠️ Descarga tus datos regularmente para evitar pérdidas")

    csv = df.to_csv(index=False).encode("utf-8")

    st.download_button("📥 Backup CSV", csv, "backup.csv")

    # Excel
    excel = BytesIO()
    df.to_excel(excel,index=False,engine='openpyxl')
    st.download_button("📊 Excel",excel.getvalue(),"data.xlsx")

    # PDF
    pdf_buffer = BytesIO()
    p = canvas.Canvas(pdf_buffer)
    y=800
    for i,row in df.iterrows():
        p.drawString(30,y,str(row.values))
        y-=20
    p.save()
    st.download_button("📄 PDF",pdf_buffer.getvalue(),"data.pdf")

    # Word
    doc=Document()
    doc.add_heading("Reporte J&A")
    for i,row in df.iterrows():
        doc.add_paragraph(str(row.values))
    word_buffer=BytesIO()
    doc.save(word_buffer)
    st.download_button("📝 Word",word_buffer.getvalue(),"data.docx")

else:
    st.info("📭 No hay datos todavía")

# ✅ RESET TOTAL
if st.button("🗑️ Reset total"):
    if os.path.exists(FILE):
        os.remove(FILE)
    st.success("✅ Datos eliminados completamente")
